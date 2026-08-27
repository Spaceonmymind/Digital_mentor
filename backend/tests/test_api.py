import asyncio
import io
from decimal import Decimal

import fitz
import pytest
from docx import Document as DocxDocument
from sqlalchemy import select

from app.api.v1 import documents as documents_api
from app.api.v1.analyses import _fallback_pdf_evidence, _locate_evidence
from app.api.v1 import internal_llm as internal_llm_api
from app.assessment.models import Assessment
from app.db.models import Analysis
from app.db.models import AnalysisEvent, AnalysisResult, LLMCall
from app.db.session import async_session_factory
from app.llm.errors import LLMConfigurationError
from app.llm.registry import AGGREGATOR, WORKER
from app.llm.schemas import LLMResult, LLMTestStructuredResponse, LLMUsage
from app.main import app
from app.methodology.models import Methodology, MethodologyCriterion, MethodologyIndicator, PromptTemplate
from app.pipeline.artifact_resolver import ArtifactResolver
from app.services import chat_service


def make_pdf_bytes(text: str = "Методы исследования\nДля анализа был использован сравнительный подход.") -> bytes:
    document = fitz.open()
    page = document.new_page()
    page.insert_text((72, 72), text)
    buffer = io.BytesIO()
    document.save(buffer)
    document.close()
    return buffer.getvalue()


def make_docx_bytes(text: str = "Методы исследования") -> bytes:
    document = DocxDocument()
    document.add_heading(text, level=2)
    document.add_paragraph("Для анализа был использован сравнительный подход.")
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def test_chat_answer_is_split_into_readable_paragraphs():
    answer = "Первое предложение. Второе предложение. Третье предложение. Четвертое предложение."
    formatted = chat_service._format_chat_answer(answer)
    assert formatted == "Первое предложение. Второе предложение.\n\nТретье предложение. Четвертое предложение."


def test_pdf_evidence_location_uses_exact_quote_and_bbox():
    payload = {
        "pages": [
            {
                "page_number": 3,
                "blocks": [
                    {"block_index": 7, "text": "Проверяемая цитата из финансовой модели", "bbox": [1, 2, 3, 4]}
                ],
            }
        ]
    }
    assert _locate_evidence(payload, "цитата из финансовой модели", None) == (3, 7, [1, 2, 3, 4], "exact")


def test_pdf_evidence_fallback_uses_real_relevant_block_and_bbox():
    payload = {
        "pages": [
            {"page_number": 1, "blocks": [{"block_index": 1, "text": "Общее описание документа без финансовых показателей и расчетов.", "bbox": [1, 2, 30, 40]}]},
            {"page_number": 4, "blocks": [{"block_index": 8, "text": "Финансовая модель содержит выручку, затраты, инвестиции и расчет окупаемости проекта.", "bbox": [10, 20, 300, 80]}]},
        ]
    }
    page, block_index, bbox, quote = _fallback_pdf_evidence(payload, "C5")
    assert (page, block_index, bbox) == (4, 8, [10, 20, 300, 80])
    assert quote.startswith("Финансовая модель")


async def upload_pdf(client):
    response = await client.post(
        "/api/v1/documents",
        files={"upload": ("work.pdf", make_pdf_bytes(), "application/pdf")},
    )
    assert response.status_code == 200, response.text
    return response.json()


async def upload_docx(client):
    response = await client.post(
        "/api/v1/documents",
        files={
            "upload": (
                "work.docx",
                make_docx_bytes("Методы исследования\nДля анализа был использован сравнительный подход."),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert response.status_code == 200, response.text
    return response.json()


async def wait_for_completed_analysis(client, analysis_id: str):
    for _ in range(50):
        status_response = await client.get(f"/api/v1/analyses/{analysis_id}")
        assert status_response.status_code == 200, status_response.text
        payload = status_response.json()
        if payload["status"] in {"completed", "failed", "cancelled"}:
            return payload
        await asyncio.sleep(0.05)
    raise AssertionError("analysis did not finish in time")


@pytest.mark.asyncio
async def test_healthcheck(client):
    response = await client.get("/health")
    assert response.status_code == 200
    assert response.json() == {"status": "ok"}

    live = await client.get("/health/live")
    assert live.status_code == 200
    assert live.json()["status"] == "live"

    ready = await client.get("/health/ready")
    assert ready.status_code == 200
    assert ready.json()["status"] == "ready"
    assert ready.json()["analysis_engine"] == "mock"


@pytest.mark.asyncio
async def test_internal_llm_test_endpoint_records_trace(client):
    class FakeLLMClient:
        async def ask(self, model: str, system_prompt: str, user_prompt: str, response_model):
            assert model == WORKER
            assert system_prompt == "Системная инструкция"
            assert "Описание идеи" in user_prompt
            output = response_model(summary="Краткое описание", keywords=["идея", "рынок"])
            assert isinstance(output, LLMTestStructuredResponse)
            return LLMResult(
                output=output,
                provider_response_id="cmpl-test",
                requested_model=model,
                actual_model="mistral-medium-3.5",
                aggregator=AGGREGATOR,
                provider="mistral",
                finish_reason="stop",
                temperature=0,
                usage=LLMUsage(
                    prompt_tokens=40,
                    completion_tokens=12,
                    total_tokens=52,
                    cached_tokens=8,
                    reasoning_tokens=0,
                    cost_rub=Decimal("0.010000"),
                ),
                latency_ms=123,
                status="success",
            )

    app.dependency_overrides[internal_llm_api.get_llm_client] = lambda: lambda: FakeLLMClient()
    try:
        response = await client.post(
            "/api/v1/internal/llm/test",
            json={"text": "Сервис для проверки гипотез", "system_prompt": "Системная инструкция"},
        )
    finally:
        app.dependency_overrides.pop(internal_llm_api.get_llm_client, None)

    assert response.status_code == 200, response.text
    assert response.json() == {
        "summary": "Краткое описание",
        "keywords": ["идея", "рынок"],
        "tokens": 52,
        "cost_rub": "0.010000",
        "provider": "mistral",
        "requested_model": WORKER,
        "actual_model": "mistral-medium-3.5",
        "provider_response_id": "cmpl-test",
        "latency_ms": 123,
    }

    async with async_session_factory() as session:
        traces = (await session.execute(select(LLMCall))).scalars().all()
    assert len(traces) == 1
    assert traces[0].model == WORKER
    assert traces[0].provider_response_id == "cmpl-test"
    assert traces[0].requested_model == WORKER
    assert traces[0].actual_model == "mistral-medium-3.5"
    assert traces[0].aggregator == AGGREGATOR
    assert traces[0].provider == "mistral"
    assert traces[0].finish_reason == "stop"
    assert traces[0].temperature == Decimal("0.00")
    assert traces[0].total_tokens == 52
    assert traces[0].cached_tokens == 8
    assert traces[0].reasoning_tokens == 0
    assert traces[0].cost_rub == Decimal("0.010000")
    assert traces[0].latency_ms == 123
    assert traces[0].status == "success"


@pytest.mark.asyncio
async def test_internal_llm_test_endpoint_records_failed_trace(client):
    class FailingLLMClient:
        async def ask(self, **_):
            raise RuntimeError("transport failure")

    app.dependency_overrides[internal_llm_api.get_llm_client] = lambda: lambda: FailingLLMClient()
    try:
        response = await client.post(
            "/api/v1/internal/llm/test",
            json={"text": "Сервис для проверки гипотез", "system_prompt": "Системная инструкция"},
        )
    finally:
        app.dependency_overrides.pop(internal_llm_api.get_llm_client, None)

    assert response.status_code == 502
    assert response.json()["error"]["code"] == "LLM_TEST_FAILED"

    async with async_session_factory() as session:
        traces = (await session.execute(select(LLMCall))).scalars().all()
    assert len(traces) == 1
    assert traces[0].requested_model == WORKER
    assert traces[0].aggregator == AGGREGATOR
    assert traces[0].status == "failed"


@pytest.mark.asyncio
async def test_internal_llm_test_endpoint_records_failed_trace_when_client_init_fails(client):
    def fail_to_create_client():
        raise LLMConfigurationError("POLZA_API_KEY is not configured")

    app.dependency_overrides[internal_llm_api.get_llm_client] = lambda: fail_to_create_client
    try:
        response = await client.post(
            "/api/v1/internal/llm/test",
            json={"text": "Сервис для проверки гипотез", "system_prompt": "Системная инструкция"},
        )
    finally:
        app.dependency_overrides.pop(internal_llm_api.get_llm_client, None)

    assert response.status_code == 500
    assert response.json()["error"]["code"] == "LLM_CONFIGURATION_ERROR"

    async with async_session_factory() as session:
        traces = (await session.execute(select(LLMCall))).scalars().all()
    assert len(traces) == 1
    assert traces[0].requested_model == WORKER
    assert traces[0].aggregator == AGGREGATOR
    assert traces[0].status == "failed"


@pytest.mark.asyncio
async def test_internal_assessment_endpoint_creates_assessment(client):
    async with async_session_factory() as session:
        methodology = Methodology(
            code="ASSESSMENT_TEST",
            name="Assessment test methodology",
            version="1.0",
            description=None,
            is_active=True,
            is_demo=True,
        )
        session.add(methodology)
        await session.commit()
        await session.refresh(methodology)

    response = await client.post(
        "/api/v1/internal/assessment",
        json={
            "artifact_type": "document",
            "artifact_id": "document-123",
            "methodology_id": methodology.id,
        },
    )

    assert response.status_code == 200, response.text
    payload = response.json()
    assert payload["id"]
    assert payload["artifact_id"] == "document-123"
    assert payload["artifact_type"] == "document"
    assert payload["methodology_id"] == methodology.id
    assert payload["status"] == "created"
    assert payload["created_at"]

    async with async_session_factory() as session:
        assessment = await session.get(Assessment, payload["id"])
    assert assessment is not None
    assert assessment.artifact_id == "document-123"
    assert assessment.methodology_id == methodology.id
    assert assessment.status == "created"


@pytest.mark.asyncio
async def test_internal_methodology_endpoint_returns_full_methodology(client):
    async with async_session_factory() as session:
        methodology = Methodology(
            code="MIRCLASS",
            name="Demo MIRCLASS methodology",
            version="1.0",
            description="Demo seed for endpoint test.",
            is_active=True,
            is_demo=True,
        )
        session.add(methodology)
        await session.flush()
        criterion = MethodologyCriterion(
            methodology_id=methodology.id,
            number="1",
            title="Demo criterion: business model",
            description="Demo criterion for endpoint test.",
            order_index=1,
            weight=Decimal("0.25"),
            is_demo=True,
        )
        session.add(criterion)
        await session.flush()
        session.add(
            MethodologyIndicator(
                criterion_id=criterion.id,
                title="Demo indicator: value proposition",
                description="Demo indicator for endpoint test.",
                expected_result="Demo expected result.",
                order_index=1,
                weight=Decimal("1.00"),
                is_demo=True,
            )
        )
        await session.commit()

    response = await client.get("/api/v1/internal/methodologies/MIRCLASS")

    assert response.status_code == 200, response.text
    payload = response.json()
    assert payload["code"] == "MIRCLASS"
    assert payload["version"] == "1.0"
    assert payload["is_demo"] is True
    assert len(payload["criteria"]) == 1
    assert payload["criteria"][0]["number"] == "1"
    assert payload["criteria"][0]["title"] == "Demo criterion: business model"
    assert payload["criteria"][0]["weight"] == "0.2500"
    assert payload["criteria"][0]["indicators"][0]["title"] == "Demo indicator: value proposition"
    assert payload["criteria"][0]["indicators"][0]["weight"] == "1.0000"
    assert payload["criteria"][0]["indicators"][0]["expected_result"] == "Demo expected result."
    assert payload["prompts"] == []


@pytest.mark.asyncio
async def test_internal_methodology_endpoint_returns_404_for_missing_methodology(client):
    response = await client.get("/api/v1/internal/methodologies/MISSING")

    assert response.status_code == 404
    assert response.json()["error"]["code"] == "METHODOLOGY_NOT_FOUND"


@pytest.mark.asyncio
async def test_internal_methodologies_endpoint_creates_and_lists_methodologies(client):
    create_response = await client.post(
        "/api/v1/internal/methodologies",
        json={
            "code": "IDEA_CARD",
            "name": "Карточка идеи",
            "description": "Тестовая методология без критериев.",
            "version": "1.0",
            "is_active": True,
        },
    )
    assert create_response.status_code == 200, create_response.text
    created = create_response.json()
    assert created["code"] == "IDEA_CARD"
    assert created["version"] == "1.0"
    assert created["is_active"] is True

    list_response = await client.get("/api/v1/internal/methodologies")
    assert list_response.status_code == 200, list_response.text
    assert [item["code"] for item in list_response.json()] == ["IDEA_CARD"]


@pytest.mark.asyncio
async def test_artifact_resolver_uses_simple_filename_rules():
    resolver = ArtifactResolver()

    assert await resolver.resolve(None, "startup-vkr.pdf", {}) == "STARTUP_VKR"
    assert await resolver.resolve(None, "работа-стартап.docx", {}) == "STARTUP_VKR"
    assert await resolver.resolve(None, "ordinary-document.pdf", {}) == "UNIVERSAL_DOCUMENT"
    assert await resolver.resolve("startup_vkr", "ordinary-document.pdf", {}) == "STARTUP_VKR"


@pytest.mark.asyncio
async def test_internal_pipeline_build_creates_assessment_and_tasks(client):
    async with async_session_factory() as session:
        methodology = Methodology(
            code="STARTUP_VKR",
            name="ВКР как стартап",
            version="1.0",
            description="Pipeline test methodology.",
            is_active=True,
            is_demo=True,
        )
        session.add(methodology)
        await session.flush()
        criterion_1 = MethodologyCriterion(
            methodology_id=methodology.id,
            number="1",
            title="Demo criterion 1",
            description="Demo criterion for pipeline test.",
            weight=Decimal("0.50"),
            order_index=1,
            is_demo=True,
        )
        criterion_2 = MethodologyCriterion(
            methodology_id=methodology.id,
            number="2",
            title="Demo criterion 2",
            description="Demo criterion for pipeline test.",
            weight=Decimal("0.50"),
            order_index=2,
            is_demo=True,
        )
        session.add_all([criterion_1, criterion_2])
        await session.flush()
        session.add_all(
            [
                MethodologyIndicator(
                    criterion_id=criterion_1.id,
                    title="Demo indicator 1",
                    description="Demo indicator for pipeline test.",
                    expected_result="Demo expected result.",
                    weight=Decimal("0.60"),
                    order_index=1,
                    is_demo=True,
                ),
                MethodologyIndicator(
                    criterion_id=criterion_1.id,
                    title="Demo indicator 2",
                    description="Demo indicator for pipeline test.",
                    expected_result="Demo expected result.",
                    weight=Decimal("0.40"),
                    order_index=2,
                    is_demo=True,
                ),
                MethodologyIndicator(
                    criterion_id=criterion_2.id,
                    title="Demo indicator 3",
                    description="Demo indicator for pipeline test.",
                    expected_result="Demo expected result.",
                    weight=Decimal("1.00"),
                    order_index=1,
                    is_demo=True,
                ),
                PromptTemplate(
                    methodology_id=methodology.id,
                    stage="worker",
                    system_prompt="Demo prompt template.",
                    user_template="Demo user template.",
                    version="1.0",
                    is_demo=True,
                ),
            ]
        )
        await session.commit()

    response = await client.post(
        "/api/v1/internal/pipeline/build",
        json={"artifact_type": "STARTUP_VKR", "artifact_id": "document-123"},
    )

    assert response.status_code == 200, response.text
    payload = response.json()
    assert payload["assessment_id"]
    assert payload["methodology"] == "STARTUP_VKR"
    assert payload["tasks_count"] == 3
    assert [task["criterion"] for task in payload["tasks"]] == [
        "Demo criterion 1",
        "Demo criterion 1",
        "Demo criterion 2",
    ]
    assert [task["indicator"] for task in payload["tasks"]] == [
        "Demo indicator 1",
        "Demo indicator 2",
        "Demo indicator 3",
    ]
    assert all(task["prompt_template_id"] for task in payload["tasks"])

    async with async_session_factory() as session:
        assessment = await session.get(Assessment, payload["assessment_id"])
    assert assessment is not None
    assert assessment.artifact_type == "STARTUP_VKR"
    assert assessment.artifact_id == "document-123"
    assert assessment.methodology_id == methodology.id


@pytest.mark.asyncio
async def test_upload_pdf_success(client):
    payload = await upload_pdf(client)
    assert payload["name"] == "work.pdf"
    assert payload["mime_type"] == "application/pdf"
    assert payload["size"] > 0
    assert payload["status"] == "uploaded"

    content_response = await client.get(f"/api/v1/documents/{payload['id']}/content")
    assert content_response.status_code == 200
    content = content_response.json()["content"]
    assert content["page_count"] == 1
    assert content["pages"][0]["blocks"]


@pytest.mark.asyncio
async def test_upload_docx_success(client):
    response = await client.post(
        "/api/v1/documents",
        files={
            "upload": (
                "work.docx",
                make_docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert response.status_code == 200, response.text
    payload = response.json()
    assert payload["name"] == "work.docx"
    assert payload["status"] == "uploaded"


@pytest.mark.asyncio
async def test_reject_wrong_format(client):
    response = await client.post(
        "/api/v1/documents",
        files={"upload": ("work.txt", b"plain text", "text/plain")},
    )
    assert response.status_code == 400
    assert response.json()["error"]["code"] == "UNSUPPORTED_FILE_TYPE"


@pytest.mark.asyncio
async def test_reject_renamed_exe_as_pdf(client):
    response = await client.post(
        "/api/v1/documents",
        files={"upload": ("malware.pdf", b"MZnot-a-pdf", "application/pdf")},
    )
    assert response.status_code == 400
    assert response.json()["error"]["code"] == "INVALID_FILE_SIGNATURE"


@pytest.mark.asyncio
async def test_reject_corrupted_pdf(client):
    response = await client.post(
        "/api/v1/documents",
        files={"upload": ("broken.pdf", b"%PDF broken", "application/pdf")},
    )
    assert response.status_code == 422
    assert response.json()["error"]["code"] == "DOCUMENT_CORRUPTED"


@pytest.mark.asyncio
async def test_reject_corrupted_docx(client):
    response = await client.post(
        "/api/v1/documents",
        files={
            "upload": (
                "broken.docx",
                b"PKbroken",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert response.status_code == 400
    assert response.json()["error"]["code"] == "INVALID_DOCX_STRUCTURE"


@pytest.mark.asyncio
async def test_reject_empty_pdf(client):
    response = await client.post(
        "/api/v1/documents",
        files={"upload": ("empty.pdf", b"", "application/pdf")},
    )
    assert response.status_code == 400
    assert response.json()["error"]["code"] == "EMPTY_FILE"


@pytest.mark.asyncio
async def test_upload_path_traversal_name_is_normalized(client):
    response = await client.post(
        "/api/v1/documents",
        files={"upload": ("../nested/work.pdf", make_pdf_bytes(), "application/pdf")},
    )
    assert response.status_code == 200, response.text
    assert response.json()["name"] == "work.pdf"


@pytest.mark.asyncio
async def test_upload_unicode_filename(client):
    response = await client.post(
        "/api/v1/documents",
        files={"upload": ("работа.pdf", make_pdf_bytes(), "application/pdf")},
    )
    assert response.status_code == 200, response.text
    assert response.json()["name"] == "работа.pdf"


@pytest.mark.asyncio
async def test_reject_file_over_size_limit(client):
    original_limit = documents_api.settings.max_upload_size_mb
    object.__setattr__(documents_api.settings, "max_upload_size_mb", 0)
    try:
        response = await client.post(
            "/api/v1/documents",
            files={"upload": ("work.pdf", make_pdf_bytes(), "application/pdf")},
        )
    finally:
        object.__setattr__(documents_api.settings, "max_upload_size_mb", original_limit)

    assert response.status_code == 413
    assert response.json()["error"]["code"] == "FILE_TOO_LARGE"


@pytest.mark.asyncio
async def test_create_analysis_progress_and_result(client):
    document = await upload_docx(client)
    create_response = await client.post(
        "/api/v1/analyses",
        json={
            "document_id": document["id"],
            "analysis_type": "mentor",
            "methodology_id": "mentor-default",
            "methodology_version": "draft",
        },
    )
    assert create_response.status_code == 200, create_response.text
    analysis_id = create_response.json()["analysis_id"]

    status = await wait_for_completed_analysis(client, analysis_id)
    assert status["status"] == "completed"
    assert status["progress"] == 100

    result_response = await client.get(f"/api/v1/analyses/{analysis_id}/result")
    assert result_response.status_code == 200, result_response.text
    result = result_response.json()
    assert result["analysis_id"] == analysis_id
    assert result["overall_score"] == 87
    assert result["criteria"]
    assert result["methodology"]["methodology_id"] == "mentor-default"
    assert "extra_blocks" in result
    assert "evidence" in result

    async with async_session_factory() as session:
        events = (
            await session.execute(select(AnalysisEvent).where(AnalysisEvent.analysis_id == analysis_id))
        ).scalars().all()
    assert len(events) >= 9
    assert max(event.progress for event in events) == 100

    report_response = await client.post(f"/api/v1/analyses/{analysis_id}/reports")
    assert report_response.status_code == 200, report_response.text
    report = report_response.json()
    pdf_response = await client.get(report["report_url"])
    assert pdf_response.status_code == 200
    assert pdf_response.content.startswith(b"%PDF")
    pdf = fitz.open(stream=pdf_response.content, filetype="pdf")
    extracted_text = "\n".join(page.get_text() for page in pdf)
    pdf.close()
    assert "Цифровой ментор" in extracted_text
    assert "Итоговый отчет" in extracted_text
    assert "Общий балл" in extracted_text

    detailed_response = await client.post(f"/api/v1/analyses/{analysis_id}/detailed-report")
    assert detailed_response.status_code == 200, detailed_response.text
    detailed_status = detailed_response.json()
    assert detailed_status["status"] in {"pending", "running", "completed"}

    for _ in range(20):
        status_response = await client.get(f"/api/v1/analyses/{analysis_id}/detailed-report/status")
        assert status_response.status_code == 200, status_response.text
        detailed_status = status_response.json()
        if detailed_status["status"] == "completed":
            break
        await asyncio.sleep(0.05)
    assert detailed_status["status"] == "completed"
    assert detailed_status["report_url"] == f"/api/v1/analyses/{analysis_id}/detailed-report/download"

    detailed_pdf_response = await client.get(detailed_status["report_url"])
    assert detailed_pdf_response.status_code == 200
    assert detailed_pdf_response.content.startswith(b"%PDF")
    detailed_pdf = fitz.open(stream=detailed_pdf_response.content, filetype="pdf")
    detailed_text = "\n".join(page.get_text() for page in detailed_pdf)
    detailed_pdf.close()
    assert "Подробный аналитический отчет" in detailed_text
    assert "Конкретные фрагменты текста" in detailed_text
    assert "Для анализа был использован сравнительный подход" in detailed_text
    assert "Почему этот фрагмент важен" in detailed_text
    assert "Что рекомендуется изменить" in detailed_text


@pytest.mark.asyncio
async def test_create_analysis_missing_document(client):
    response = await client.post(
        "/api/v1/analyses",
        json={
            "document_id": "missing",
            "analysis_type": "mentor",
            "methodology_id": "mentor-default",
            "methodology_version": "draft",
        },
    )
    assert response.status_code == 404
    assert response.json()["error"]["code"] == "DOCUMENT_NOT_FOUND"


@pytest.mark.asyncio
async def test_analysis_history_metrics_and_source_are_read_only(client):
    document = await upload_pdf(client)
    create_response = await client.post(
        "/api/v1/analyses",
        json={
            "document_id": document["id"],
            "analysis_type": "mentor",
            "methodology_id": "STARTUP_VKR",
            "methodology_version": "2.0",
            "mode": "demo",
        },
    )
    analysis_id = create_response.json()["analysis_id"]
    await wait_for_completed_analysis(client, analysis_id)
    async with async_session_factory() as session:
        session.add(
            LLMCall(
                model="openai/gpt-test",
                requested_model="openai/gpt-test",
                actual_model="gpt-test",
                aggregator="polza",
                provider="openai",
                analysis_id=analysis_id,
                agent_code="A-01",
                prompt_tokens=100,
                completion_tokens=40,
                total_tokens=140,
                cached_tokens=10,
                cost_rub=Decimal("1.250000"),
                latency_ms=900,
                status="success",
            )
        )
        await session.commit()

    history = await client.get("/api/v1/analyses/history?limit=10&offset=0")
    assert history.status_code == 200, history.text
    history_payload = history.json()
    assert history_payload["total"] == 1
    assert history_payload["items"][0]["analysis_id"] == analysis_id
    assert history_payload["items"][0]["document_name"] == "work.pdf"
    assert history_payload["items"][0]["overall_score"] == 87

    metrics = await client.get(f"/api/v1/analyses/{analysis_id}/metrics")
    assert metrics.status_code == 200, metrics.text
    metrics_payload = metrics.json()
    assert metrics_payload["llm_calls_count"] == 1
    assert metrics_payload["total_tokens"] == 140
    assert metrics_payload["cost_rub"] == "1.250000"
    assert metrics_payload["agents"][0]["agent_code"] == "A-01"
    serialized = metrics.text.lower()
    assert "prompt" not in serialized
    assert "api_key" not in serialized

    source = await client.get(f"/api/v1/documents/{document['id']}/source")
    assert source.status_code == 200
    assert source.content.startswith(b"%PDF")
    assert source.headers["content-disposition"].startswith("inline")

    preview = await client.get(f"/api/v1/documents/{document['id']}/pages/1/preview")
    assert preview.status_code == 200
    assert preview.headers["content-type"] == "image/png"
    assert preview.content.startswith(b"\x89PNG")

    evidence = await client.get(f"/api/v1/analyses/{analysis_id}/evidence")
    assert evidence.status_code == 200
    assert evidence.json() == []


@pytest.mark.asyncio
async def test_cancel_analysis(client):
    document = await upload_pdf(client)
    async with async_session_factory() as session:
        analysis = Analysis(
            document_id=document["id"],
            analysis_type="mentor",
            methodology_id="mentor-default",
            methodology_version="draft",
            status="queued",
            progress=0,
            current_step="queued",
        )
        session.add(analysis)
        await session.commit()
        await session.refresh(analysis)
        analysis_id = analysis.id

    response = await client.post(f"/api/v1/analyses/{analysis_id}/cancel")
    assert response.status_code == 200
    assert response.json()["status"] == "cancelled"


@pytest.mark.asyncio
async def test_delete_document(client):
    document = await upload_pdf(client)
    response = await client.delete(f"/api/v1/documents/{document['id']}")
    assert response.status_code == 200, response.text
    assert response.json()["status"] == "deleted"

    get_response = await client.get(f"/api/v1/documents/{document['id']}")
    assert get_response.status_code == 404


@pytest.mark.asyncio
async def test_tts_falls_back_without_provider(client):
    response = await client.post("/api/v1/tts", json={"text": "Анализ завершен", "voice_id": "mentor-default"})
    assert response.status_code == 200, response.text
    payload = response.json()
    assert payload["status"] == "fallback"
    assert payload["format"] == "mp3"
    assert payload["provider"] == "browser"
    assert payload["source"] == "browser"
    assert payload["audio_url"] is None


@pytest.mark.asyncio
async def test_chat_without_analysis_id_is_rejected(client):
    response = await client.post("/api/v1/chat/messages", json={"message": "Почему снижена оценка?"})
    assert response.status_code == 422


@pytest.mark.asyncio
async def test_chat_response_for_analysis(client):
    document = await upload_pdf(client)
    create_response = await client.post(
        "/api/v1/analyses",
        json={
            "document_id": document["id"],
            "analysis_type": "mentor",
            "methodology_id": "mentor-default",
            "methodology_version": "draft",
        },
    )
    analysis_id = create_response.json()["analysis_id"]
    await wait_for_completed_analysis(client, analysis_id)

    response = await client.post(
        "/api/v1/chat/messages",
        json={"analysis_id": analysis_id, "message": "Как получить больше 90 баллов?"},
    )
    assert response.status_code == 200, response.text
    payload = response.json()
    assert payload["message_id"]
    assert "90" in payload["answer"]


@pytest.mark.asyncio
async def test_startup_vkr_chat_uses_relevant_document_fragments(client, monkeypatch):
    captured = {}

    class FakeChatLLM:
        async def ask(self, model, system_prompt, user_prompt, response_model, **kwargs):
            captured["user_prompt"] = user_prompt
            captured["system_prompt"] = system_prompt
            captured["max_completion_tokens"] = kwargs.get("max_completion_tokens")
            return LLMResult(
                output=response_model(answer="Покажите фрагмент про сравнительный подход и добавьте критерии проверки."),
                provider_response_id="chat-test",
                requested_model=model,
                actual_model=model,
                aggregator=AGGREGATOR,
                provider="fake",
                finish_reason="stop",
                usage=LLMUsage(prompt_tokens=10, completion_tokens=5, total_tokens=15, cost_rub=Decimal("0.01")),
                latency_ms=1,
                status="success",
            )

    monkeypatch.setattr(chat_service, "LLMClient", lambda: FakeChatLLM())
    document = await upload_docx(client)
    create_response = await client.post(
        "/api/v1/analyses",
        json={
            "document_id": document["id"],
            "analysis_type": "mentor",
            "methodology_id": "mentor-default",
            "methodology_version": "draft",
        },
    )
    analysis_id = create_response.json()["analysis_id"]
    await wait_for_completed_analysis(client, analysis_id)
    async with async_session_factory() as session:
        analysis = await session.get(Analysis, analysis_id)
        analysis.methodology_id = "STARTUP_VKR"
        result = (
            await session.execute(select(AnalysisResult).where(AnalysisResult.analysis_id == analysis_id).limit(1))
        ).scalar_one()
        result.result_json = {
            **result.result_json,
            "remarks": [{"title": "Недостаточно раскрыт сравнительный подход", "recommendation": "Добавить критерии проверки."}],
        }
        session.add_all([analysis, result])
        await session.commit()

    response = await client.post(
        "/api/v1/chat/messages",
        json={"analysis_id": analysis_id, "message": "Какой фрагмент текста мне править?"},
    )

    assert response.status_code == 200, response.text
    assert "Релевантные фрагменты исходного документа" in captured["user_prompt"]
    assert "Для анализа был использован сравнительный подход" in captured["user_prompt"]
    assert captured["max_completion_tokens"] == 1400
    assert "простыми словами" in captured["system_prompt"]
    assert "Разделяй ответ пустыми строками" in captured["system_prompt"]
