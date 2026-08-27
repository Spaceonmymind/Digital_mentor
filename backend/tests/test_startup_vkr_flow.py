import io
from decimal import Decimal

import pytest
from docx import Document as DocxDocument
from sqlalchemy import select

from app.assessment.models import IndicatorResult
from app.db.models import Analysis, AnalysisResult, Document, LLMCall
from app.db.session import async_session_factory
from app.execution.models import AgentResult, AgentTaskRun, GateDecision, MentorAnalysisResult
from app.execution.schemas import CriticOutput, DemoAgentOutput, DemoFinalReport, MentorReport, WorkerIndicatorOutput
from app.execution.startup_vkr import StartupVkrAgentFlow
from app.llm.registry import CRITIC, FINAL_EXPERT, WORKER
from app.llm.errors import LLMResponseValidationError
from app.llm.schemas import LLMResult, LLMUsage
from app.methodology.models import Methodology, MethodologyAgent, MethodologyCriterion
from app.methodology.seeds.startup_vkr.data import A15_CHECKS
from app.methodology.seeds.startup_vkr import ensure_startup_vkr_seed
from app.pipeline.schemas import PipelineBuildRequest
from app.pipeline.service import PipelineService
from app.services.extraction import TextExtractionService
from app.services.reports import ReportService
from app.services.storage import DocumentStorage


class FakeUpload:
    def __init__(self, content: bytes):
        self.content = content
        self.done = False

    async def read(self, _size: int = -1):
        if self.done:
            return b""
        self.done = True
        return self.content


def docx_bytes(text: str) -> bytes:
    document = DocxDocument()
    document.add_paragraph(text)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


async def seed_document(text: str):
    from app.db.models import Document

    storage = DocumentStorage()
    stored = await storage.save(FakeUpload(docx_bytes(text)), ".docx")
    document = Document(
        original_name="startup-vkr-demo.docx",
        stored_name=stored.stored_name,
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        size=stored.size,
        checksum=stored.checksum,
        storage_path=str(stored.path),
        extraction_status="completed",
        status="uploaded",
    )
    async with async_session_factory() as session:
        session.add(document)
        await session.commit()
        await session.refresh(document)
        extracted_path = storage.extracted_path(document.id)
        TextExtractionService().extract(document.id, stored.path, ".docx", extracted_path)
        document.extracted_path = str(extracted_path)
        await session.commit()
        await session.refresh(document)
        return document


class FakeStartupLLM:
    def __init__(self):
        self.calls = []

    def _mentor_report(self) -> MentorReport:
        return MentorReport(
            header={
                "work_title": "Demo startup VKR",
                "work_type": "ВКР как стартап",
                "analysis_date": "2026-08-07",
                "work_version": None,
                "methodology": "STARTUP_VKR 1.1",
                "current_stage": "S2",
            },
            what_this_work_is="Работа предлагает цифровой сервис, но механизм результата и экономика принятия пока описаны частично.",
            veto={
                "is_active": False,
                "reason": None,
                "why_further_assessment_is_meaningless": None,
                "how_to_remove": None,
            },
            what_survived=["Идея цифрового сервиса сохраняется как предмет дальнейшей проверки."],
            objections=[
                {
                    "title": "Механизм результата пока не доказан",
                    "what_does_not_work": "Из описания не видно, какой компонент производит заявленный эффект.",
                    "why": "Без цепочки действий нельзя отличить работающий механизм от названия технологии.",
                    "where_to_move": "Описать участников, данные, действия и проверяемый результат каждого шага.",
                }
            ],
            one_question={"question": "Какой механизм делает заявленный эффект воспроизводимым?"},
            one_next_step={
                "step": "Постройте схему действий сервиса от входных данных до проверяемого результата.",
                "check_result": "По схеме можно указать, кто выполняет действие и чем подтверждается результат.",
            },
            stage_assessments=[
                {
                    "stage_code": "S1",
                    "title": "Проблема",
                    "score": 3,
                    "completed": "Проблема заявлена и частично описана.",
                    "next_level_requirement": "Показать наблюдаемое событие и масштаб.",
                },
                {
                    "stage_code": "S2",
                    "title": "Противоречие",
                    "score": 2,
                    "completed": "Противоречие намечено.",
                    "next_level_requirement": "Разделить требование результата и ограничение, которое ему мешает.",
                },
            ],
            mentor_block={
                "leading_blind_spot": "Автор смешивает эффект и механизм.",
                "what_changed": "Главный фокус смещен с оформления на проверку конструкции.",
                "what_remains_unresolved": "Не доказана воспроизводимость результата.",
                "mentor_question": "Где в работе можно увидеть механизм, а не обещание?",
                "recommended_intervention": "Попросить автора нарисовать цепочку действий и проверить каждое звено.",
            },
            spoken_summary=(
                "Я закончил разбор. У работы есть исходная идея цифрового сервиса. "
                "Главное противоречие сейчас в том, что механизм результата пока не доказан. "
                "Ответьте сначала на один вопрос: какой механизм делает заявленный эффект воспроизводимым. "
                "Следующий шаг: построить схему действий сервиса от входных данных до проверяемого результата."
            ),
        )

    async def ask(self, model, system_prompt, user_prompt, response_model, **kwargs):
        self.calls.append({"model": model, "response_model": response_model.__name__, "system_prompt": system_prompt, "user_prompt": user_prompt, **kwargs})
        if response_model is WorkerIndicatorOutput:
            output = WorkerIndicatorOutput(
                status="partially_satisfied",
                score=55,
                summary="Механизм частично описан.",
                strengths=["Есть описание решения"],
                issues=["Есть неподтвержденный вывод"],
                evidence=[{"quote": "цитата которой нет", "page": None, "section": None, "explanation": "Worker сослался на фрагмент"}],
                recommendations=["Уточнить механизм"],
                confidence=0.7,
            )
        elif response_model is CriticOutput:
            output = CriticOutput(
                verdict="revise",
                worker_result_supported=False,
                confirmed_findings=[],
                disputed_findings=[{"worker_finding": "Механизм частично описан", "reason": "Цитата не подтверждена", "evidence": []}],
                missed_issues=["Не проверена экономика принятия"],
                contradictions=[],
                recommended_adjustments=["Попросить автора подтвердить ключевые утверждения"],
                confidence=0.8,
            )
        elif response_model is DemoAgentOutput:
            assigned = {
                "C1, C3": ["C1", "C3"],
                "C2, C6": ["C2", "C6"],
                "C4, C5": ["C4", "C5"],
                "C3, C6": ["C3", "C6"],
            }
            codes = next((value for marker, value in assigned.items() if marker in user_prompt), ["C1"])
            output = DemoAgentOutput(
                summary="Demo-агент нашел ключевой вывод.",
                criteria=[
                    {
                        "criterion_code": code,
                        "score_recommendation": 8,
                        "summary": "Критерий раскрыт, но доказательства нужно уточнить.",
                        "strengths": ["Есть предмет проверки"],
                        "issues": ["Нужно уточнить доказательства"],
                        "evidence": [{"section": "Резюме", "quote": None}],
                        "confidence": 0.8,
                    }
                    for code in codes
                ],
                recommendations=["Сжать аргументацию до проверяемой схемы"],
            )
        elif response_model is DemoFinalReport:
            output = DemoFinalReport(
                overall_score=47,
                criteria=[
                    {"code": "C1", "name": "Проблема и актуальность", "score": 8, "comment": "Проблема показана достаточно ясно.", "strengths": ["Потребность описана"], "issues": ["Нужны источники"]},
                    {"code": "C2", "name": "Инновационность и продукт", "score": 8, "comment": "Продукт понятен, но механизм нужно уточнить.", "strengths": ["Продукт определен"], "issues": ["MVP описан кратко"]},
                    {"code": "C3", "name": "Рынок и целевая аудитория", "score": 7, "comment": "Сегмент назван, гипотезы подтверждены частично.", "strengths": ["Сегмент выделен"], "issues": ["Мало интервью"]},
                    {"code": "C4", "name": "Бизнес-модель", "score": 8, "comment": "Модель дохода описана.", "strengths": ["Плательщик понятен"], "issues": ["Каналы требуют проверки"]},
                    {"code": "C5", "name": "Финансовая реализуемость", "score": 7, "comment": "Финансы есть, предпосылки требуют проверки.", "strengths": ["Затраты оценены"], "issues": ["Нет сценариев"]},
                    {"code": "C6", "name": "Риски и развитие", "score": 9, "comment": "Риски и roadmap раскрыты.", "strengths": ["Есть roadmap"], "issues": ["Нужны владельцы рисков"]},
                ],
                strengths=["Есть актуальная проблема", "Есть архитектурная идея", "Есть экономическое обоснование"],
                remarks=["Не хватает отказных сценариев", "Не хватает go/no-go порогов", "Нужно уточнить доказательства"],
                recommendations=["Описать fallback", "Добавить пороги решения", "Сжать механизм в схему"],
                conclusion="Demo-разбор показывает сильную основу и несколько проверяемых направлений доработки.",
                spoken_summary="Demo-анализ завершен: общий балл 47 из 60, главный следующий шаг — уточнить fallback и пороги.",
                disclaimer="Предварительная аналитическая оценка документа не заменяет решение ГЭК.",
            )
        else:
            output = self._mentor_report()
        return LLMResult(
            output=output,
            provider_response_id=f"fake-{len(self.calls)}",
            requested_model=model,
            actual_model=model,
            aggregator="polza.ai",
            provider="fake-provider",
            finish_reason="stop",
            usage=LLMUsage(prompt_tokens=10, completion_tokens=20, total_tokens=30, cost_rub=Decimal("0.01")),
            latency_ms=10,
            status="success",
        )


class FailingDemoFinalLLM(FakeStartupLLM):
    async def ask(self, model, system_prompt, user_prompt, response_model, **kwargs):
        if response_model is DemoFinalReport:
            self.calls.append({"model": model, "response_model": response_model.__name__, "user_prompt": user_prompt, **kwargs})
            raise RuntimeError("truncated demo final json")
        return await super().ask(model, system_prompt, user_prompt, response_model, **kwargs)


class EmptyA28DemoLLM(FakeStartupLLM):
    async def ask(self, model, system_prompt, user_prompt, response_model, **kwargs):
        if response_model is DemoAgentOutput and "C3, C6" in user_prompt:
            self.calls.append({"model": model, "response_model": response_model.__name__, "user_prompt": user_prompt, **kwargs})
            raise LLMResponseValidationError({"reason": "empty content"})
        return await super().ask(model, system_prompt, user_prompt, response_model, **kwargs)


@pytest.mark.asyncio
async def test_startup_vkr_seed_creates_real_methodology_without_demo_plan_items():
    document = await seed_document("ВКР описывает стартап и механизм цифрового сервиса.")
    async with async_session_factory() as session:
        methodology = await ensure_startup_vkr_seed(session)
        criteria = (await session.execute(select(MethodologyCriterion).where(MethodologyCriterion.methodology_id == methodology.id))).scalars().all()
        agents = (await session.execute(select(MethodologyAgent).where(MethodologyAgent.methodology_id == methodology.id))).scalars().all()
        pipeline = await PipelineService(session).build(
            PipelineBuildRequest(artifact_type="STARTUP_VKR", artifact_id=document.id, filename=document.original_name)
        )

    assert methodology.is_demo is False
    assert methodology.version == "2.0"
    assert {criterion.number for criterion in criteria} == {"C1", "C2", "C3", "C4", "C5", "C6"}
    assert {criterion.source for criterion in criteria} == {"FinUniversity VKR Startup Regulation, Order №3136/o dated 20.12.2023"}
    assert all(criterion.weight is None for criterion in criteria)
    assert len(agents) == 9
    assert pipeline.tasks_count == 18
    assert not any(task.criterion.startswith("Demo") for task in pipeline.tasks)


@pytest.mark.asyncio
async def test_startup_vkr_keeps_old_version_and_uses_new_active_version():
    async with async_session_factory() as session:
        session.add(
            Methodology(
                id="startup-vkr-old-version",
                code="STARTUP_VKR",
                name="ВКР как стартап",
                description="Old reproducible version",
                version="1.0",
                is_active=True,
                is_demo=False,
                source="anti_during_methodology",
            )
        )
        await session.commit()
        methodology = await ensure_startup_vkr_seed(session)

        versions = (
            await session.execute(select(Methodology.version).where(Methodology.code == "STARTUP_VKR").order_by(Methodology.version))
        ).scalars().all()
        active_versions = (
            await session.execute(
                select(Methodology.version).where(Methodology.code == "STARTUP_VKR", Methodology.is_active.is_(True))
            )
        ).scalars().all()

    assert "1.0" in versions
    assert "1.1" in versions
    assert "2.0" in versions
    assert methodology.version == "2.0"
    assert active_versions == ["2.0"]


def test_a15_revision_126_contains_nine_checks():
    assert len(A15_CHECKS) == 9
    assert "заявленное свойство противоречит устройству" in A15_CHECKS


@pytest.mark.asyncio
async def test_startup_vkr_flow_runs_worker_critic_and_one_final_expert():
    document = await seed_document("Автор предлагает цифровой сервис. Механизм описан частично. Экономика требует проверки.")
    llm = FakeStartupLLM()
    async with async_session_factory() as session:
        await ensure_startup_vkr_seed(session)
        pipeline = await PipelineService(session).build(
            PipelineBuildRequest(artifact_type="STARTUP_VKR", artifact_id=document.id, filename=document.original_name)
        )
        result = await StartupVkrAgentFlow(session, llm_client=llm).execute(pipeline.assessment_id, analysis_id="analysis-1")

    assert result.status == "requires_revision"
    assert result.methodology_version == "2.0"
    assert result.total_tokens == 690
    assert result.report is not None
    assert result.report.header.current_stage == "S2"
    assert len([result.report.one_question.question]) == 1
    assert result.report.one_next_step.step
    assert all(0 <= item.score <= 5 for item in result.report.stage_assessments)
    assert all(item.next_level_requirement for item in result.report.stage_assessments)
    assert result.technical is not None
    report_text = result.report.model_dump_json()
    assert "Worker" not in report_text
    assert "Critic" not in report_text
    assert "quote_not_found" not in report_text
    assert [call["model"] for call in llm.calls].count(WORKER) == 18
    assert [call["model"] for call in llm.calls].count(CRITIC) == 4
    assert [call["model"] for call in llm.calls].count(FINAL_EXPERT) == 1

    async with async_session_factory() as session:
        indicator_results = (await session.execute(select(IndicatorResult))).scalars().all()
        agent_runs = (await session.execute(select(AgentTaskRun))).scalars().all()
        agent_results = (await session.execute(select(AgentResult))).scalars().all()
        gates = (await session.execute(select(GateDecision))).scalars().all()
        llm_calls = (await session.execute(select(LLMCall))).scalars().all()
        mentor_results = (await session.execute(select(MentorAnalysisResult))).scalars().all()

    assert all((item.evidence_json[0]["quote"] is None) for item in indicator_results)
    assert len(agent_runs) == 5
    assert len(agent_results) == 5
    assert {gate.decision_source for gate in gates} == {"demo_auto_approve"}
    assert {call.requested_model for call in llm_calls} == {WORKER, CRITIC, FINAL_EXPERT}
    assert len(mentor_results) == 1

    async with async_session_factory() as session:
        public_report = await StartupVkrAgentFlow(session).student_result(mentor_results[0].assessment_id)
        technical = await StartupVkrAgentFlow(session).technical_result(mentor_results[0].assessment_id)

    public_payload = public_report
    assert "mentor_block" not in public_payload
    assert technical.total_tokens == 690


@pytest.mark.asyncio
async def test_startup_vkr_flow_cache_hit_does_not_call_llm_again():
    document = await seed_document("Повторный запуск должен использовать сохраненные результаты.")
    llm = FakeStartupLLM()
    async with async_session_factory() as session:
        await ensure_startup_vkr_seed(session)
        pipeline = await PipelineService(session).build(
            PipelineBuildRequest(artifact_type="STARTUP_VKR", artifact_id=document.id, filename=document.original_name)
        )
        await StartupVkrAgentFlow(session, llm_client=llm).execute(pipeline.assessment_id, analysis_id="analysis-1")
        await StartupVkrAgentFlow(session, llm_client=llm).execute(pipeline.assessment_id, analysis_id="analysis-1")

    assert len(llm.calls) == 23


@pytest.mark.asyncio
async def test_startup_vkr_demo_flow_keeps_agents_and_returns_short_scored_report():
    document = await seed_document("Резюме проекта. Проблема KYC. Архитектура DID Wallet Verifier. Экономика SAM NPV. Риски отказов.")
    llm = FakeStartupLLM()
    async with async_session_factory() as session:
        await ensure_startup_vkr_seed(session)
        pipeline = await PipelineService(session).build(
            PipelineBuildRequest(artifact_type="STARTUP_VKR", artifact_id=document.id, filename=document.original_name)
        )
        payload, metrics = await StartupVkrAgentFlow(session, llm_client=llm).execute_demo(pipeline.assessment_id, analysis_id="analysis-demo")
        stored_demo = (
            await session.execute(
                select(MentorAnalysisResult).where(MentorAnalysisResult.assessment_id == pipeline.assessment_id)
            )
        ).scalar_one()

    assert payload.overall_score == 47
    assert stored_demo.analysis_id == "analysis-demo"
    assert stored_demo.document_id == document.id
    assert stored_demo.status == "completed"
    assert stored_demo.result_json["extra_blocks"]["assessment_id"] == pipeline.assessment_id
    assert payload.extra_blocks["mode"] == "demo"
    assert payload.extra_blocks["demo_report"]["overall_score"] == 47
    assert [call["model"] for call in llm.calls].count(WORKER) == 2
    assert [call["model"] for call in llm.calls].count(CRITIC) == 2
    assert [call["model"] for call in llm.calls].count(FINAL_EXPERT) == 1
    assert {call["response_model"] for call in llm.calls} == {"DemoAgentOutput", "DemoFinalReport"}
    assert {call["max_completion_tokens"] for call in llm.calls if call["response_model"] == "DemoAgentOutput"} == {1200}
    assert {call["max_completion_tokens"] for call in llm.calls if call["response_model"] == "DemoFinalReport"} == {1800}
    demo_agent_prompts = [call["system_prompt"] for call in llm.calls if call["response_model"] == "DemoAgentOutput"]
    assert all("учебную работу" in prompt and "Не требуй" in prompt for prompt in demo_agent_prompts)
    final_prompt = next(call["system_prompt"] for call in llm.calls if call["response_model"] == "DemoFinalReport")
    assert "не применяй к ним второй штраф" in final_prompt
    assert "простыми словами" in final_prompt
    assert metrics["mode"] == "demo"
    assert metrics["agent_time_a15"] >= 0
    assert payload.methodology.methodology_version == "2.0"
    assert [item.code for item in payload.criteria] == ["C1", "C2", "C3", "C4", "C5", "C6"]
    assert payload.overall_score == sum(item.score for item in payload.criteria)

    analysis = Analysis(
        id="demo-report-analysis",
        document_id=document.id,
        analysis_type="mentor",
        methodology_id="STARTUP_VKR",
        methodology_version="2.0",
        status="completed",
        progress=100,
    )
    result = AnalysisResult(analysis_id=analysis.id, result_json=payload.model_dump(mode="json"))
    pdf_text = "\n".join(ReportService()._build_lines(analysis, document, result.result_json))
    assert "Предварительная оценка документа ВКР-стартапа" in pdf_text
    assert "C1. Проблема и актуальность" in pdf_text
    assert "не заменяет решение ГЭК" in pdf_text
    assert "Assessment ID" not in pdf_text
    assert "Токены" not in pdf_text
    assert ReportService()._render_pdf(pdf_text.splitlines()).startswith(b"%PDF")


@pytest.mark.asyncio
async def test_startup_vkr_demo_flow_falls_back_when_final_json_is_invalid():
    document = await seed_document("Резюме проекта. Проблема KYC. Архитектура DID Wallet Verifier. Экономика SAM NPV. Риски отказов.")
    llm = FailingDemoFinalLLM()
    async with async_session_factory() as session:
        await ensure_startup_vkr_seed(session)
        pipeline = await PipelineService(session).build(
            PipelineBuildRequest(artifact_type="STARTUP_VKR", artifact_id=document.id, filename=document.original_name)
        )
        payload, _ = await StartupVkrAgentFlow(session, llm_client=llm).execute_demo(pipeline.assessment_id, analysis_id="analysis-demo-fallback")
        task_run = (
            await session.execute(
                select(AgentTaskRun).where(
                    AgentTaskRun.assessment_id == pipeline.assessment_id,
                    AgentTaskRun.model_role == "final_expert",
                )
            )
        ).scalar_one()
        final_result = (
            await session.execute(
                select(AgentResult).where(
                    AgentResult.assessment_id == pipeline.assessment_id,
                    AgentResult.model_role == "final_expert",
                )
            )
        ).scalar_one()

    assert payload.extra_blocks["mode"] == "demo"
    assert payload.overall_score is not None
    assert task_run.status == "completed"
    assert task_run.error_code == "DEMO_FINAL_FALLBACK"
    assert final_result.llm_call_id is None
    assert "собран из результатов независимых проверок" in payload.verdict


@pytest.mark.asyncio
async def test_startup_vkr_demo_retries_and_isolates_empty_agent_response():
    document = await seed_document("Резюме проекта. Проблема KYC. Архитектура DID Wallet Verifier. Экономика SAM NPV. Риски отказов.")
    llm = EmptyA28DemoLLM()
    async with async_session_factory() as session:
        await ensure_startup_vkr_seed(session)
        pipeline = await PipelineService(session).build(
            PipelineBuildRequest(artifact_type="STARTUP_VKR", artifact_id=document.id, filename=document.original_name)
        )
        payload, metrics = await StartupVkrAgentFlow(session, llm_client=llm).execute_demo(
            pipeline.assessment_id,
            analysis_id="analysis-demo-agent-fallback",
        )
        runs = (
            await session.execute(
                select(AgentTaskRun).where(AgentTaskRun.assessment_id == pipeline.assessment_id)
            )
        ).scalars().all()
        failed_calls = (
            await session.execute(
                select(LLMCall).where(
                    LLMCall.assessment_id == pipeline.assessment_id,
                    LLMCall.status == "failed",
                )
            )
        ).scalars().all()

    a28_calls = [call for call in llm.calls if call["response_model"] == "DemoAgentOutput" and "C3, C6" in call["user_prompt"]]
    a28_run = next(run for run in runs if run.agent_code == "A-28")
    assert len(a28_calls) == 2
    assert payload.extra_blocks["mode"] == "demo"
    assert payload.overall_score is not None
    assert metrics["agent_time_a28"] >= 0
    assert a28_run.status == "completed"
    assert a28_run.error_code == "DEMO_AGENT_FALLBACK"
    assert len(failed_calls) == 1
    assert failed_calls[0].agent_code == "A-28"
    assert failed_calls[0].requested_model == CRITIC
    assert failed_calls[0].max_completion_tokens == 1200


def test_mentor_report_rejects_internal_terms_and_invalid_stage_score():
    valid = {
        "header": {
            "work_title": "Demo",
            "work_type": "ВКР как стартап",
            "analysis_date": "2026-08-07",
            "work_version": None,
            "methodology": "STARTUP_VKR 1.1",
            "current_stage": "S2",
        },
        "what_this_work_is": "Короткое описание работы.",
        "veto": {"is_active": False, "reason": None, "why_further_assessment_is_meaningless": None, "how_to_remove": None},
        "what_survived": ["Есть проверяемый предмет обсуждения."],
        "objections": [
            {
                "title": "Главное возражение",
                "what_does_not_work": "Не показан механизм.",
                "why": "Без механизма невозможно проверить заявленный эффект.",
                "where_to_move": "Описать цепочку действий.",
            }
        ],
        "one_question": {"question": "Как механизм создает результат?"},
        "one_next_step": {"step": "Построить схему механизма.", "check_result": "Каждый шаг имеет вход и выход."},
        "stage_assessments": [
            {
                "stage_code": "S2",
                "title": "Противоречие",
                "score": 5,
                "completed": "Стадия описана.",
                "next_level_requirement": "Проверить воспроизводимость.",
            }
        ],
        "mentor_block": {
            "leading_blind_spot": "Смешаны эффект и механизм.",
            "what_changed": "Фокус уточнен.",
            "what_remains_unresolved": "Нужна проверка механизма.",
            "mentor_question": "Где проверяемая цепочка?",
            "recommended_intervention": "Попросить автора показать схему.",
        },
        "spoken_summary": "Я закончил разбор. Главное противоречие связано с механизмом. Следующий шаг — построить схему.",
    }
    MentorReport.model_validate(valid)
    domain_text = dict(valid)
    domain_text["what_this_work_is"] = "Работа обсуждает токены доступа, стоимость внедрения и экономические издержки как предмет анализа."
    MentorReport.model_validate(domain_text)
    invalid = dict(valid)
    invalid["what_this_work_is"] = "Worker сообщил Assessment ID 12345678-1234-1234-1234-123456789abc."
    with pytest.raises(ValueError):
        MentorReport.model_validate(invalid)
    invalid_score = dict(valid)
    invalid_score["stage_assessments"] = [dict(valid["stage_assessments"][0], score=100)]
    with pytest.raises(ValueError):
        MentorReport.model_validate(invalid_score)


def test_demo_schemas_do_not_emit_unsupported_array_size_keywords():
    agent_schema = DemoAgentOutput.model_json_schema()
    final_schema = DemoFinalReport.model_json_schema()

    schema_text = f"{agent_schema}{final_schema}"

    for keyword in ("maxItems", "minItems", "maximum", "minimum", "maxLength", "minLength"):
        assert keyword not in schema_text


def test_demo_scores_accept_percent_like_model_output():
    agent = DemoAgentOutput.model_validate(
        {
            "summary": "Короткий вывод.",
            "criteria": [{
                "criterion_code": "C1",
                "score_recommendation": 55,
                "summary": "Проблема раскрыта частично.",
                "strengths": ["Есть проблема"],
                "issues": ["Мало доказательств"],
                "evidence": [{"section": "Введение", "quote": None}],
                "confidence": 0.7,
            }],
            "recommendations": ["Уточнить расчет"],
        }
    )
    assert agent.criteria[0].score_recommendation == 6

    report = DemoFinalReport.model_validate(
        {
            "overall_score": 55,
            "criteria": [
                {"code": "C1", "name": "Проблема и актуальность", "score": "80%", "comment": "Проблема видна.", "strengths": [], "issues": []},
                {"code": "C2", "name": "Инновационность и продукт", "score": 70, "comment": "Продукт описан.", "strengths": [], "issues": []},
                {"code": "C3", "name": "Рынок и целевая аудитория", "score": "6/10", "comment": "Рынок намечен.", "strengths": [], "issues": []},
                {"code": "C4", "name": "Бизнес-модель", "score": 50, "comment": "Модель требует проверки.", "strengths": [], "issues": []},
                {"code": "C5", "name": "Финансовая реализуемость", "score": 4, "comment": "Финансы неполны.", "strengths": [], "issues": []},
                {"code": "C6", "name": "Риски и развитие", "score": 90, "comment": "Риски заявлены.", "strengths": [], "issues": []},
            ],
            "strengths": ["Сильное ядро"],
            "remarks": ["Нужна проверка"],
            "recommendations": ["Собрать доказательства"],
            "conclusion": "Работу можно быстро улучшить.",
            "spoken_summary": "Я закончил разбор.",
            "disclaimer": "Предварительная оценка не заменяет решение ГЭК.",
        }
    )
    assert [item.score for item in report.criteria] == [8, 7, 6, 5, 4, 9]
    assert report.overall_score == 39


def test_student_pdf_uses_mentor_report_without_technical_dump():
    report = FakeStartupLLM()._mentor_report().model_dump(mode="json")
    analysis = Analysis(
        id="analysis-without-uuid-like-value",
        document_id="document-1",
        analysis_type="mentor",
        methodology_id="STARTUP_VKR",
        methodology_version="1.1",
        status="completed",
        progress=100,
    )
    document = Document(
        id="document-1",
        original_name="demo.docx",
        stored_name="demo.docx",
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        size=1,
        checksum="checksum",
        storage_path="/tmp/demo.docx",
        extraction_status="completed",
        status="uploaded",
    )
    result = AnalysisResult(analysis_id=analysis.id, result_json={"extra_blocks": {"mentor_report": report}})
    lines = ReportService()._build_lines(analysis, document, result.result_json)
    text = "\n".join(lines)

    assert "ЦИФРОВОЙ МЕНТОР" in text
    assert "Assessment ID" not in text
    assert "Токены" not in text
    assert "Стоимость" not in text
    assert "Worker" not in text
    assert "Critic" not in text
